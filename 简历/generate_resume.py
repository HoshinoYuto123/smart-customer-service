"""Generate a concise one-page Chinese resume."""
from docx import Document
import os
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup (tight margins for one-page fit) ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.0)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

# Default style
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.25
# Set East-Asian font
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), '宋体')
rPr.insert(0, rFonts)

# ── Helpers ──
def run_font(run, name, size, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    run_font(r, '黑体', size, True)
    return p

def add_info(text, size=9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    run_font(r, '宋体', size)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_section(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    b = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '2E75B6'})
    pBdr.append(b); pPr.append(pBdr)
    r = p.add_run(text)
    run_font(r, '黑体', 12, True, (0x2E, 0x75, 0xB6))

def add_bullet(text, size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('• ' + text)
    run_font(r, '宋体', size)

def add_sub(text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    run_font(r, '黑体', size, True)

def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    rId = paragraph.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    h = paragraph._p.makeelement(qn('w:hyperlink'), {
        qn('r:id'): rId,
        qn('w:history'): '1',
    })
    r = h.makeelement(qn('w:r'), {})
    rPr_elem = r.makeelement(qn('w:rPr'), {})
    rFonts = rPr_elem.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr_elem.append(rFonts)
    c = rPr_elem.makeelement(qn('w:color'), {qn('w:val'): '0563C1'})
    rPr_elem.append(c)
    u = rPr_elem.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    rPr_elem.append(u)
    sz = rPr_elem.makeelement(qn('w:sz'), {qn('w:val'): '18'})
    rPr_elem.append(sz)
    r.append(rPr_elem)
    t = r.makeelement(qn('w:t'), {qn('xml:space'): 'preserve'})
    t.text = text
    r.append(t)
    h.append(r)
    paragraph._p.append(h)

def add_sub_info(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    run_font(r, '宋体', 9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ═══════════════════════════════════════
# HEADER (table layout: info left, photo right)
# ═══════════════════════════════════════
header_table = doc.add_table(rows=1, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Remove table borders
for row in header_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = '宋体'
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
        for edge in ('top', 'left', 'bottom', 'right'):
            element = tcBorders.makeelement(qn(f'w:{edge}'), {qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0'})
            tcBorders.append(element)
        tcPr.append(tcBorders)

# Left cell: personal info
left_cell = header_table.cell(0, 0)
left_cell.width = Cm(12)

p = left_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(2)
r = p.add_run('方 智 强')
run_font(r, '黑体', 20, True)

p2 = left_cell.add_paragraph()
p2.paragraph_format.space_after = Pt(1)
r2 = p2.add_run('电话：13155982667  |  邮箱：13155982667@163.com  |  现居：江西南昌')
run_font(r2, '宋体', 9)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p3 = left_cell.add_paragraph()
r3a = p3.add_run('江西财经大学 · 软件工程 · 本科（2027届）  |  GitHub：')
run_font(r3a, '宋体', 9)
r3a.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
add_hyperlink(p3, 'github.com/HoshinoYuto123', 'https://github.com/HoshinoYuto123')

# Right cell: photo
right_cell = header_table.cell(0, 1)
right_cell.width = Cm(4.5)
p_photo = right_cell.paragraphs[0]
p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
photo_path = os.path.join(os.path.dirname(__file__), 'photo.jpg')
if os.path.exists(photo_path):
    r_photo = p_photo.add_run()
    r_photo.add_picture(photo_path, width=Inches(1.1), height=Inches(1.45))

# ═══════════════════════════════════════
# 求职意向
# ═══════════════════════════════════════
add_section('求职意向')
add_bullet('目标岗位：AI 应用开发 / 后端开发工程师（实习生）')
add_bullet('独立完成从需求到上线的 AI 项目（智能客服 Agent），全程使用 VibeCoding（Claude Code）辅助开发，了解 LangGraph + FastAPI + RAG 基础技术栈。')

# ═══════════════════════════════════════
# 核心技能
# ═══════════════════════════════════════
add_section('核心技能')
add_bullet('编程语言：掌握 C++ 基础（语法、STL、面向对象），了解 Python 基本使用。')
add_bullet('AI 应用：了解 LangChain / LangGraph 框架基础，理解 Agent 多节点状态机工作流程；了解 Prompt Engineering 与 Function Calling 工具调用机制。')
add_bullet('RAG 检索增强生成：了解向量数据库（ChromaDB）基本使用、FAQ 知识库构建与混合检索（向量 + 关键词）基础流程。')
add_bullet('工具与部署：有 Office 办公软件使用经验；了解 FastAPI 基础、Git 版本控制、Docker 容器化部署；熟练使用 Claude Code 等 AI 工具进行辅助开发。')

# ═══════════════════════════════════════
# 项目经历
# ═══════════════════════════════════════
add_section('核心项目经历')

add_sub('SCS-Agent — 智能客服 Agent 系统（VibeCoding 实践）')
# Project experience sub-info with hyperlink
p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(1)
r_sub1 = p_sub.add_run('技术栈：Python, LangGraph, FastAPI, ChromaDB, DeepSeek API  |  GitHub：')
run_font(r_sub1, '宋体', 9)
r_sub1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
add_hyperlink(p_sub, 'github.com/HoshinoYuto123/smart-customer-service', 'https://github.com/HoshinoYuto123/smart-customer-service')
r_sub2 = p_sub.add_run('  |  已部署上线')
run_font(r_sub2, '宋体', 9)
r_sub2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
add_bullet('使用 Claude Code 进行 VibeCoding 辅助开发，通过自然语言描述需求、拆解任务，由 AI 协助完成代码生成与调试，完成从 0 到 1 的项目构建。')
add_bullet('基于 LangChain / LangGraph 框架搭建多节点 Agent 状态机（澄清 → 路由 → 检索执行 → 回复 → 降级），通过 Prompt 工程引导大模型行为，结合 Function Calling 实现工具调用。')
add_bullet('构建 FAQ 知识库（4 个业务域 × 40 条问答）与物流信息数据库（12 条记录），实现基础的 RAG 检索问答，支持向量语义搜索与关键词匹配。')
add_bullet('设计电商风格聊天前端界面，通过 SQLite 持久化存储实现历史会话管理与上下文记忆，支持对话恢复与 Token 感知的上下文截断。')
add_bullet('部署至 Render 平台实现公网访问，项目已开源（95+ 文件、38 个单元测试）。')

# ═══════════════════════════════════════
# 教育背景
# ═══════════════════════════════════════
add_section('教育背景')
add_sub_info('江西财经大学  |  软件工程  |  本科（2023.09 - 2027.07）')
add_bullet('主修：软件工程、数据结构、计算机网络、操作系统、计算机组成原理（408 考纲全覆盖），具备扎实计算机基础与系统工程思维。')

# ═══════════════════════════════════════
# 校园经历
# ═══════════════════════════════════════
add_section('校园经历')
add_sub_info('校学生会组织部 · 干事（2023.09 - 2024.06）')
add_bullet('对接全院各班级团支书，定期跟进入党入团情况摸排，将收集的信息整理归档并汇总成表，确保材料规范、按时上报。')
add_bullet('参与微信公众号"团干培训"相关推文的图文编辑与排版，积累了一定的新媒体内容制作经验。')

# ═══════════════════════════════════════
# 自我评价
# ═══════════════════════════════════════
add_section('自我评价')
add_bullet('善于利用 AI 工具（Claude Code）进行 VibeCoding，能将想法快速转化为可运行的产品原型，注重实践与快速落地。')
add_bullet('有从 0 到 1 完成项目的经验，具备基本的需求拆解、代码组织与文档撰写能力，做事踏实、不浮躁。')
add_bullet('对 AI / LLM 领域有浓厚兴趣，持续关注行业动态，愿意在实际工作中深入学习与成长。')

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), '方智强_求职简历_智能客服Agent.docx')
doc.save(output_path)
print(f'Done: {output_path}')
